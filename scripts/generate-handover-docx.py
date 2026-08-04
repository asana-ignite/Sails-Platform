#!/usr/bin/env python3
"""Generate SAILS Platform Handover Document (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── Helper Functions ──
def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row_obj in table.rows:
                row_obj.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SAILS Platform")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Technical Handover Document")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8C)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}").font.size = Pt(12)
meta.add_run("\n")
meta.add_run("Audience: Solution Architect / Technical Lead").font.size = Pt(12)
meta.add_run("\n")
meta.add_run("Classification: Internal — Confidential").font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    "Multi-tenant, schema-per-tenant internal operating system platform.\n"
    "Monorepo: packages/core (Next.js API), packages/console (Vite React UI), packages/shared (TypeScript contracts)."
)
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Architecture Overview",
    "3. Technology Stack",
    "4. Environment & Configuration",
    "5. Database Architecture",
    "6. Security Pipeline",
    "7. API Reference",
    "8. Frontend Architecture (Console)",
    "9. Zoning Multi-Tenancy Architecture",
    "10. Deployment & DevOps",
    "11. Backup & Recovery Procedures",
    "12. Known Issues, Golden Rules & Operational Playbook",
    "13. Roadmap & Technical Debt",
    "Appendix A: Key File Reference",
    "Appendix B: Key Commands Reference",
    "Appendix C: Environment Variables Reference",
    "Appendix D: Prisma Model Inventory",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "SAILS (Ignite Idea Operating System) is a high-performance, multi-tenant internal operating system platform "
    "built with Bun, TypeScript, and PostgreSQL. It operates as a headless backend API (SAILS Core) with a "
    "decoupled Progressive Web Application frontend (SAILS Console). The system enables Ignite Idea to define "
    "custom data structures (Sales Leads, Project Tasks, Cases, Timesheets) which are dynamically translated "
    "into native PostgreSQL tables in real-time."
)

doc.add_heading("Current Status", level=2)
doc.add_paragraph("The platform has completed Phases 1–7 (Foundation, Multi-Tenancy, Security, UI, Advanced Metadata). Phase 8 (Internal Modules) is in progress. The production deployment pipeline (CI/CD, cloud infrastructure) has NOT been configured yet — this is the first priority for the incoming Solution Architect.")

doc.add_heading("Key Statistics", level=2)
add_styled_table(doc,
    ["Metric", "Value"],
    [
        ["Monorepo packages", "3 (core, console, shared)"],
        ["Prisma models", "21 (in core schema) + 3 (in global schema, planned)"],
        ["API endpoints", "35+"],
        ["Database migrations", "12 applied"],
        ["Field type plugins", "20 (19 core + shared registry)"],
        ["UI plugin registries", "4 (Admin, FieldControl, Action, Widget)"],
        ["System capabilities", "15 (across 5 categories)"],
        ["Current phase", "Phase 8 (Internal Modules) — In Progress"],
        ["Development DB volume", "klaoplatform_pgdata (external — DO NOT DELETE)"],
        ["Development login", "admin@klao.app / Welcome2Ignite"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. Architecture Overview", level=1)

doc.add_heading("2.1 Monorepo Structure", level=2)
doc.add_paragraph("The project is a Bun workspace monorepo with three packages:")
add_styled_table(doc,
    ["Package", "Path", "Runtime", "Purpose"],
    [
        ["sails-core", "packages/core", "Bun / Next.js 14", "Headless API engine — auth, metadata, dynamic CRUD, Prisma ORM"],
        ["@sails/console", "packages/console", "Bun / Vite 5 / React 18", "SPA admin UI — plugin-based, DB-driven navigation, PWA-ready"],
        ["@sails/shared", "packages/shared", "TypeScript (no runtime)", "Shared type contracts, field type registry, validation, permissions"],
    ],
    [3, 3.5, 3, 6.5]
)

doc.add_heading("2.2 System Topology", level=2)
doc.add_paragraph("The platform runs as three Docker containers orchestrated by a single docker-compose.yml:")
add_bullet(doc, "sails-db: PostgreSQL 16 Alpine (port 5433 → 5432)")
add_bullet(doc, "sails-core: Next.js API server on Bun (port 3000)")
add_bullet(doc, "sails-console: Vite dev server for React SPA (port 5173)")

doc.add_paragraph(
    "Data flow: Browser → Console (Vite, port 5173) → API Proxy /api/* → Core (Next.js, port 3000) → "
    "PostgreSQL (port 5432). The Vite dev server proxies all /api/* requests to the Core container via "
    "Docker internal networking."
)

doc.add_heading("2.3 Architectural Principles", level=2)
add_bullet(doc, "Schema-Per-Tenant Isolation: Each tenant gets a dedicated PostgreSQL schema (tenant_{name}) with its own RLS policies. Metadata, auth, and audit tables reside in the shared core schema.")
add_bullet(doc, "Metadata-Driven DDL: Table definitions stored in core.tables/core.fields are materialized into physical PostgreSQL tables by AlchemaCore at provisioning time. No DDL is hand-written for tenant data.")
add_bullet(doc, "Navigation is DB-Driven: core.console_apps + core.console_menus → GET /api/console/config → ConsoleContext → Sidebar.tsx. Zero hardcoded menu items at runtime.")
add_bullet(doc, "Plugin Architecture: Field types, UI controls, admin pages, and widgets all use registry-based resolution with React.lazy() code splitting.")
add_bullet(doc, "Two-Layer Security: Application-level RBAC (AccessGuard) + Database-level RLS (TransactionContext → SET LOCAL + PostgreSQL RLS policies).")
add_bullet(doc, "Non-Blocking Audit: All logging (data changes, system events, DDL operations) uses fire-and-forget async writes to avoid impacting request latency.")
add_bullet(doc, "CUID Primary Keys: All IDs use CUID (time-ordered strings), never auto-increment integers or UUIDv4. Migrated from UUID in migration 20260719061531.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. Technology Stack", level=1)

doc.add_heading("3.1 Core Package (sails-core)", level=2)
add_styled_table(doc,
    ["Category", "Technology", "Version", "Purpose"],
    [
        ["Runtime", "Bun", "latest", "Native TypeScript execution, binary compilation"],
        ["Framework", "Next.js", "14.x", "API routes (App Router), headless — no UI"],
        ["ORM", "Prisma", "6.19.x", "Core schema metadata management"],
        ["Query Builder", "pg + pg-format", "8.x / 1.x", "Dynamic DDL/DML generation (injection-proof)"],
        ["Auth", "next-auth", "4.24.x", "JWT strategy, Credentials + Google OAuth providers"],
        ["Password Hashing", "bcryptjs", "3.x", "12 salt rounds"],
        ["Validation", "Zod", "3.25.x", "Dynamic schema generation from field metadata"],
        ["Auth Adapter", "@auth/prisma-adapter", "2.11.x", "NextAuth ↔ Prisma user/account/session mapping"],
    ],
    [3, 3.5, 2, 7.5]
)

doc.add_heading("3.2 Console Package (@sails/console)", level=2)
add_styled_table(doc,
    ["Category", "Technology", "Version", "Purpose"],
    [
        ["Bundler", "Vite", "5.2.x", "Dev server + production build"],
        ["UI Framework", "React", "18.2.x", "SPA with react-router-dom v6"],
        ["Rich Text", "Tiptap", "3.29.x", "ProseMirror-based editor with 12 extensions"],
        ["Icons", "lucide-react", "0.378.x", "Exclusive icon library — no external icon fonts"],
        ["Color Picker", "react-colorful", "5.8.x", "Inline color picker for theme customization"],
        ["CSS", "Vanilla CSS + BEM", "—", "No frameworks (no Tailwind, no Bootstrap). sails- prefix. CSS custom properties."],
        ["Font", "Lexend", "—", "Google Fonts, weight 100–900"],
    ],
    [3, 3.5, 2, 7.5]
)

doc.add_heading("3.3 Shared Package (@sails/shared)", level=2)
doc.add_paragraph("Pure TypeScript library — no runtime dependencies, no build step. Consumed directly by both core and console via Bun workspace resolution. Contains all API contract types, field type metadata registry, system permission registry, validation utilities, and date/time formatting helpers.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. ENVIRONMENT & CONFIGURATION
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. Environment & Configuration", level=1)

doc.add_heading("4.1 Docker Compose Services", level=2)
add_styled_table(doc,
    ["Service", "Container Name", "Image", "Port Mapping"],
    [
        ["PostgreSQL", "sails-db", "postgres:16-alpine", "5433:5432 (host:container)"],
        ["Core API", "sails-core", "sails-core:latest (build)", "3000:3000"],
        ["Console UI", "sails-console", "sails-console:latest (build)", "5173:5173"],
    ],
    [3.5, 3.5, 5, 4]
)

doc.add_heading("4.2 Critical Docker Volume", level=2)
add_bold_para(doc, "WARNING: The PostgreSQL data volume is EXTERNAL and MUST NEVER BE DELETED:")
add_code_block(doc, "volumes:\n  pgdata:\n    external: true\n    name: klaoplatform_pgdata")
doc.add_paragraph("This volume contains all live tenant data. It persists across container rebuilds. Accidental deletion means total data loss.")

doc.add_heading("4.3 Environment Variables", level=2)
doc.add_paragraph("Only one .env file exists: packages/core/.env (checked into git). All other variables come from docker-compose.yml.")

add_styled_table(doc,
    ["Variable", "Value", "Source", "Purpose"],
    [
        ["DATABASE_URL", "postgresql://postgres:mysecretpassword@db:5432/postgres?schema=core", ".env + compose", "Primary DB connection (Prisma + pg Pool)"],
        ["DEFAULT_TENANT_ID", "cmrxlaeys001iky2dlttomtrw", ".env", "Fallback tenant for unauthenticated config requests"],
        ["NEXTAUTH_SECRET", "development-secret-only", "compose", "JWT signing secret (CHANGE in production)"],
        ["NEXTAUTH_URL", "http://localhost:5173", "compose", "NextAuth callback URL base"],
        ["GOOGLE_CLIENT_ID", "${GOOGLE_CLIENT_ID:-your-google-client-id}", "compose", "Google OAuth client (placeholder)"],
        ["GOOGLE_CLIENT_SECRET", "${GOOGLE_CLIENT_SECRET:-your-google-client-secret}", "compose", "Google OAuth secret (placeholder)"],
        ["VITE_CORE_URL", "http://localhost:3000", "compose (console)", "Console → Core API URL"],
        ["PLATFORM_MODE", "(unset — standalone)", "—", "When 'zoned', enables multi-database routing"],
        ["ZONE_ID", "(unset — zone-01 default)", "—", "Zone identifier in zoned mode"],
        ["LOG_DATABASE_URL", "(optional)", "—", "Separate DB for audit logs"],
        ["TEST_SESSION_JSON", "(used by tests)", "—", "CLI/test session injection (no browser needed)"],
    ],
    [3.5, 5.5, 2.5, 4.5]
)

doc.add_heading("4.4 Database Connection Details", level=2)
add_styled_table(doc,
    ["Context", "Connection String"],
    [
        ["In-container (Docker DNS)", "postgresql://postgres:mysecretpassword@db:5432/postgres?schema=core"],
        ["From host (port mapping)", "postgresql://postgres:mysecretpassword@localhost:5433/postgres?schema=core"],
        ["Production (PgBouncer)", "Auto-adds ?pgbouncer=true&connection_limit=20 when NODE_ENV=production"],
    ],
    [4, 12]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. DATABASE ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Database Architecture", level=1)

doc.add_heading("5.1 Schema Segregation", level=2)
doc.add_paragraph("The platform uses PostgreSQL schemas for logical data isolation:")

add_styled_table(doc,
    ["Schema", "Purpose", "Management", "Access"],
    [
        ["core", "Metadata, auth, audit, console config", "Prisma-managed (21 models)", "API only — no direct staff access"],
        ["tenant_{name}", "Per-tenant dynamic data tables", "AlchemaCore DDL engine", "RLS-enforced per request"],
        ["global (planned)", "Zoning control plane registry", "Prisma (global.prisma)", "Super Admin War Room only"],
    ],
    [3, 5, 4, 4]
)

doc.add_heading("5.2 Core Schema Models (21 tables)", level=2)
add_styled_table(doc,
    ["Model", "Table", "Purpose"],
    [
        ["Tenant", "tenants", "Multi-tenant registry; schemaName is unique per tenant"],
        ["TableDefinition", "tables", "Dynamic data model definitions (Leads, Customers, etc.)"],
        ["FieldDefinition", "fields", "Column metadata (physicalType + logicalType)"],
        ["ValidationRule", "validation_rules", "Field-level CHECK constraint rules"],
        ["TableLayout", "table_layouts", "UI view configurations (draft/active lifecycle)"],
        ["ConsoleApp", "console_apps", "App switcher entries (Sales, Marketing, etc.)"],
        ["ConsoleMenu", "console_menus", "Sidebar navigation tree (self-referencing hierarchy)"],
        ["ConsoleWidget", "console_widgets", "Widget bar items (Quick Accept, Agent Chat)"],
        ["Team", "teams", "Hierarchical RBAC groups (self-referencing parent_id)"],
        ["ObjectPermission", "object_permissions", "Granular CRUD on objects (per Team/Position/User)"],
        ["SystemPermission", "system_permissions", "Functional capability grants per team"],
        ["User", "users", "Core identity (NextAuth + SAILS context)"],
        ["UserTeam", "user_teams", "M:N user-team membership"],
        ["Account", "accounts", "OAuth/OIDC linked provider accounts"],
        ["Session", "sessions", "NextAuth DB sessions (unused — JWT strategy active)"],
        ["VerificationToken", "verification_tokens", "Email verification tokens"],
        ["Position", "positions", "Job position templates with prefix-based naming"],
        ["PositionSlot", "position_slots", "Individual position assignments (head count)"],
        ["TeamPosition", "team_positions", "M:N team-position assignments"],
        ["CompanyProfile", "company_profiles", "Tenant-wide branding/security/localization settings"],
        ["DataAuditLog / SystemEventLog / DdlLog", "data_audit_logs / system_event_logs / ddl_logs", "3-tier audit trail"],
    ],
    [4, 3.5, 8.5]
)

doc.add_heading("5.3 Dynamic Table Columns (Auto-Generated)", level=2)
doc.add_paragraph("Every tenant data table automatically gets these standard columns upon creation:")
add_code_block(doc,
    "id              VARCHAR(30) PRIMARY KEY    -- time-ordered CUID\n"
    "tenant_id       VARCHAR(30) NOT NULL       -- from app.current_tenant_id\n"
    "created_at      TIMESTAMPTZ DEFAULT NOW()\n"
    "updated_at      TIMESTAMPTZ DEFAULT NOW()\n"
    "owner_id        VARCHAR(30) NOT NULL       -- from app.current_user_id\n"
    "owner_team_id   VARCHAR(30)                -- from app.current_team_id\n"
    "created_by      VARCHAR(30)\n"
    "updated_by      VARCHAR(30)"
)

doc.add_heading("5.4 Migrations (12 Applied)", level=2)
add_styled_table(doc,
    ["#", "Migration Name", "Date", "Summary"],
    [
        ["1", "auth_schema_core", "2026-05-02", "Initial schema: tenants, tables, fields, users, audit"],
        ["2", "add_google_fields", "2026-05-12", "Added teams, console_apps, console_menus, Google OAuth fields"],
        ["3", "add_user_title", "2026-05-12", "Added title column to users"],
        ["4", "init_cuid_schema", "2026-07-19", "Converted all PKs from UUID to TEXT (CUID format)"],
        ["5", "logging_restructure", "2026-07-19", "Split audit_logs into 3 tables (data, system, DDL)"],
        ["6", "add_field_description", "2026-07-21", "Added description column to fields"],
        ["7", "add_layout_system_name_description", "2026-07-28", "Added system_name + description to table_layouts"],
        ["8", "add_app_config_fields", "2026-07-29", "Added slug + description to console_apps"],
        ["9", "add_data_model_to_menus", "2026-07-30", "Added data_model_id FK on console_menus"],
        ["10", "add_layout_draft_active", "2026-07-30", "Added draft/active status + published_config to layouts"],
        ["11", "add_list_view_to_menus", "2026-07-30", "Added list_view_id to console_menus"],
        ["12", "layout_system_name_global_unique + menu_path_unique", "2026-08-02", "Global unique constraints on system_name and menu path"],
    ],
    [0.8, 5.5, 2, 7.7]
)

doc.add_heading("5.5 Tenant Provisioning Flow", level=2)
doc.add_paragraph("TenantProvisioner.ts executes the following steps atomically:")
add_bullet(doc, "1. normalizeSchemaName(name) → generates snake_case schema name (e.g., tenant_acme_corp)")
add_bullet(doc, "2. generateUniqueSchemaName() → handles collisions by appending _1, _2, etc.")
add_bullet(doc, "3. engine.createTenantSchema(schema) → CREATE SCHEMA + GRANT privileges + RLS grants")
add_bullet(doc, "4. db.tenant.create() → Core metadata record + 'System Administrator' team")
add_bullet(doc, "5. Create admin user → TENANT_ADMIN role")
add_bullet(doc, "6. provisionSystemApps() → 'Settings & Admin' with full menu tree (18 sub-menus)")
add_bullet(doc, "7. provisionBusinessApps() → CRM, Sales, Dashboard apps")
add_bullet(doc, "8. provisionDefaultWidgets() → Quick Accept, Agent Chat")
add_bullet(doc, "9. provisionStandardDataModels() → Leads, Customers, Companies, Orders, Invoices (creates physical tables with RLS)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. SECURITY PIPELINE
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Security Pipeline", level=1)

doc.add_heading("6.1 5-Step Pipeline (Strict Order)", level=2)
doc.add_paragraph("Every data request must pass through all five stages. Failure at any stage results in an immediate HTTP error.")

add_styled_table(doc,
    ["Step", "Component", "File", "What It Does"],
    [
        ["1. Auth", "getAppSession()", "src/lib/auth/session.ts", "Resolves JWT identity via NextAuth. Falls back to TEST_SESSION_JSON for CLI/tests. AsyncLocalStorage caching."],
        ["2. RBAC", "AccessGuard", "src/core/engine/AccessGuard.ts", "Checks ObjectPermission (CRUD) + SystemPermission (capabilities). SUPER_ADMIN/TENANT_ADMIN fast-path bypass."],
        ["3. RLS", "TransactionContext", "src/core/engine/TransactionContext.ts", "SET ROLE rls_user + set_config(app.current_*) to activate PostgreSQL RLS policies."],
        ["4. DML+Audit", "QueryLayer", "src/core/engine/QueryLayer.ts", "Atomic CRUD + async audit log (fire-and-forget, non-blocking)."],
        ["5. Verify", "test-security.ts", "src/tests/test-security.ts", "8 RLS integration test scenarios, all passing."],
    ],
    [1.5, 3, 4.5, 7]
)

doc.add_heading("6.2 Authentication Details", level=2)
add_bullet(doc, "Strategy: JWT (not database sessions)")
add_bullet(doc, "Credentials Provider: email + bcrypt password (12 salt rounds)")
add_bullet(doc, "Google OAuth Provider: domain-restricted to @igniteidea.ai and @ignite-idea.com")
add_bullet(doc, "JIT Provisioning: Google users without existing User records are auto-created with MEMBER role on first sign-in")
add_bullet(doc, "Inactive Check: Users with isActive: false are rejected at sign-in")
add_bullet(doc, "JWT Cache: 60-second in-memory cache (jwtCache) avoids repeated DB queries on token refresh")
add_bullet(doc, "Session Methods: getAppSession() (optional), requireSession() (authenticated guard), requireAdmin() (admin guard)")

doc.add_heading("6.3 RLS Policy (Generated Per Table)", level=2)
doc.add_paragraph("Every dynamic table gets ENABLE ROW LEVEL SECURITY + FORCE ROW LEVEL SECURITY + this policy:")
add_code_block(doc,
    "CREATE POLICY {table}_owner_policy ON {schema}.{table}\n"
    "FOR ALL\n"
    "USING (\n"
    "  tenant_id = current_setting('app.current_tenant_id', true)\n"
    "  AND (\n"
    "    -- Record owner\n"
    "    owner_id = current_setting('app.current_user_id', true)\n"
    "    -- Team ownership\n"
    "    OR owner_team_id = current_setting('app.current_team_id', true)\n"
    "    -- Explicit object permission grants\n"
    "    OR EXISTS (SELECT 1 FROM core.object_permissions p\n"
    "      WHERE p.object_name = '{table_name}'\n"
    "      AND (p.team_id IN user_teams\n"
    "        OR p.user_id = current_user_id\n"
    "        OR p.position_id IN user_positions)\n"
    "      AND (p.read_scope = 'TEAM' OR 'HIERARCHY'))\n"
    "  )\n"
    ")"
)

doc.add_paragraph("For SUPER_ADMIN users, role is set to undefined (not rls_user), effectively bypassing all RLS policies. Performance indexes are created on tenant_id, owner_id, owner_team_id, and created_at for every table.")

doc.add_heading("6.4 System Permission Capabilities (15 Total)", level=2)
add_styled_table(doc,
    ["Category", "Capabilities"],
    [
        ["IAM", "system.users.manage, system.teams.manage, system.roles.assign"],
        ["Platform", "system.schema.manage, system.apps.manage, system.menus.manage, system.workflow.manage"],
        ["Security", "system.security.sso, system.security.tokens"],
        ["Extensions", "system.security.apps, system.extensions.byoc, system.integrations.api"],
        ["Operations", "system.settings.profile, system.settings.edit, system.audit.view, system.billing.manage"],
    ],
    [3, 13]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. API REFERENCE
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. API Reference", level=1)

doc.add_heading("7.1 Authentication", level=2)
add_styled_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET/POST", "/api/auth/[...nextauth]", "NextAuth handler (sign in, sign out, session, CSRF)"],
        ["POST", "/api/auth/register", "Email/password registration"],
    ],
    [2.5, 7, 6.5]
)

doc.add_heading("7.2 Console UI Configuration (9 endpoints)", level=2)
add_styled_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/api/console/config", "THE navigation endpoint — apps + menus + widgets, role-filtered"],
        ["GET/POST", "/api/console/apps", "List/create console apps"],
        ["GET/PUT/DELETE", "/api/console/apps/[id]", "Single app CRUD"],
        ["GET/POST", "/api/console/menus", "List/create menu items"],
        ["GET/PUT/DELETE", "/api/console/menus/[id]", "Single menu CRUD"],
        ["GET/POST", "/api/console/layouts", "Table layout configurations"],
        ["GET", "/api/console/permissions", "System permissions registry"],
        ["GET/PUT", "/api/console/company-profile", "Company branding/settings"],
        ["GET", "/api/console/audit-logs", "Audit log viewer"],
        ["GET/POST", "/api/console/widgets", "Widget bar items"],
    ],
    [2.5, 7, 6.5]
)

doc.add_heading("7.3 Dynamic Data", level=2)
add_styled_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/api/dynamic/[tableName]", "List records (RLS-enforced, paginated, filterable)"],
        ["POST", "/api/dynamic/[tableName]", "Create record (atomic audit logging)"],
        ["PATCH", "/api/dynamic/[tableName]", "Update record (captures before/after snapshot)"],
        ["DELETE", "/api/dynamic/[tableName]", "Delete record (captures old values for audit)"],
    ],
    [2.5, 7, 6.5]
)

doc.add_heading("7.4 Metadata Management", level=2)
add_styled_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET/POST", "/api/metadata/objects", "List/create table definitions (admin-only)"],
        ["GET/PUT/DELETE", "/api/metadata/objects/[id]", "Single table definition CRUD"],
        ["POST", "/api/metadata/fields", "Create field definition"],
        ["GET/PUT/DELETE", "/api/metadata/fields/[id]", "Single field definition CRUD"],
        ["POST", "/api/metadata/fields/[id]/reset-sequence", "Reset auto-number sequence"],
        ["GET", "/api/metadata/[tableName]", "Table metadata lookup"],
        ["GET", "/api/metadata/field-types", "Field type registry"],
    ],
    [2.5, 7, 6.5]
)

doc.add_heading("7.5 Tenant Management (14 endpoints)", level=2)
add_styled_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/api/tenant/provision", "Provision new tenant (name + adminEmail/existingUserId)"],
        ["GET/POST", "/api/tenant/users", "List/create users in tenant"],
        ["GET/PUT/DELETE", "/api/tenant/users/[id]", "Single user CRUD"],
        ["GET/POST", "/api/tenant/teams", "List/create teams"],
        ["GET/PUT/DELETE", "/api/tenant/teams/[id]", "Single team CRUD"],
        ["GET", "/api/tenant/teams/[id]/members", "Team members list"],
        ["DELETE", "/api/tenant/teams/[id]/members/[userId]", "Remove member from team"],
        ["GET", "/api/tenant/teams/[id]/object-permissions", "Team object permissions"],
        ["GET", "/api/tenant/teams/[id]/positions", "Team position assignments"],
        ["GET/POST", "/api/tenant/positions", "List/create position templates"],
        ["GET/PUT/DELETE", "/api/tenant/positions/[id]", "Single position CRUD"],
        ["GET", "/api/tenant/positions/[id]/object-permissions", "Position object permissions"],
        ["DELETE", "/api/tenant/positions/slots/[slotId]", "Remove position slot"],
    ],
    [2.5, 7, 6.5]
)

doc.add_heading("7.6 User & Zone", level=2)
add_styled_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/api/users/me", "Current user profile + granular permissions"],
        ["GET", "/api/zone/health", "Zone telemetry (CPU, memory, DB status, tenant count)"],
    ],
    [2.5, 7, 6.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. FRONTEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Frontend Architecture (Console)", level=1)

doc.add_heading("8.1 Provider Hierarchy", level=2)
add_code_block(doc,
    "BrowserRouter\n"
    "  ThemeProvider     (dynamic CSS variables from company profile)\n"
    "    AuthProvider    (NextAuth session + login/logout)\n"
    "      <Routes>\n"
    "        Public: /login, /admin-login\n"
    "        Protected:\n"
    "          ConsoleProvider  (apps, menus, widgets, header state)\n"
    "            AppLayout      (Topbar + Sidebar + WidgetBar + Mobile)\n"
    "              SmartPageRouter → DynamicTablePage / DynamicDetailPage / AppPluginShell"
)

doc.add_heading("8.2 Four Plugin Registries", level=2)

doc.add_paragraph("AdminPluginRegistry (src/features/admin/registry.tsx):")
doc.add_paragraph("Maps componentKey strings (from core.console_menus) to React.lazy() components. 18 registered plugins: AdminCompanyProfile, AdminUserManager, AdminEntityManager, AdminMenuManager, AdminPermissions, LayoutStudio, etc.")

doc.add_paragraph("FieldControlRegistry (src/features/controls/FieldControlRegistry.ts):")
doc.add_paragraph("Singleton mapping logical field types to edit/display React components. 19 built-in controls: ShortText, LongText, RichText, Number, Decimal, Currency, Percent, Select, Boolean (3 variants), Date, Time, DateTime, User, Lookup, Address, Attachment, LatLng, AutoNumber.")

doc.add_paragraph("ActionRegistry (src/features/actions/ActionRegistry.ts):")
doc.add_paragraph("Singleton for toolbar-level actions on list views. Currently one action: CreateAction (navigates to detail form). Extensible with custom action plugins.")

doc.add_paragraph("WidgetRegistry (src/features/widgets/registry.tsx):")
doc.add_paragraph("Maps widget componentKey to lazy components: OmniChannelQuickAccept, AgentChatWindows.")

doc.add_heading("8.3 Navigation System (DB-Driven)", level=2)
doc.add_paragraph("The entire navigation tree comes from the database — zero hardcoded menu items at runtime:")
add_code_block(doc,
    "core.console_apps + core.console_menus\n"
    "    ↓\n"
    "GET /api/console/config    (role-filtered by requiredCapability)\n"
    "    ↓\n"
    "ConsoleContext.navigationItems     (cached 30s)\n"
    "    ↓\n"
    "Sidebar.tsx               (recursive tree, collapsible, flyout submenus)"
)
doc.add_paragraph("Key routing: /{appSlug}/{menuPath} (e.g., /crm/leads). Legacy flat paths like /admin/company-profile still supported. Mock data fallback in getMockData() indicates the DB query returned 0 rows — NOT a code bug.")

doc.add_heading("8.4 Key Pages", level=2)
add_styled_table(doc,
    ["Page", "File", "Purpose"],
    [
        ["DynamicTablePage", "pages/DynamicTablePage.tsx", "Universal list view: fetches layout config, renders rows, pagination, sorting, filtering, actions"],
        ["DynamicDetailPage", "pages/DynamicDetailPage.tsx", "Universal record detail: new/edit/read modes, section/tab/field blocks, validation"],
        ["AppPluginShell", "pages/admin/AppPluginShell.tsx", "Resolves componentKey → AdminPluginRegistry → lazy component"],
        ["Login", "pages/Login.tsx", "Password + SSO login (Google, Microsoft, SAML toggled by config)"],
        ["UserManager", "pages/custom/UserManager.tsx", "Full user CRUD management"],
        ["ObjectManager", "pages/custom/ObjectManager.tsx", "Entity/data model management"],
        ["LayoutStudio", "pages/custom/LayoutStudio.tsx", "Layout editor for list/detail/form views"],
    ],
    [4, 5.5, 6.5]
)

doc.add_heading("8.5 Design System", level=2)
add_bullet(doc, "Methodology: BEM CSS, sails- prefix on all classes")
add_bullet(doc, "No frameworks: No Tailwind, no Bootstrap — pure vanilla CSS")
add_bullet(doc, "Theme: CSS custom properties in design-tokens.css, toggled via data-sails-theme='light' | 'dark'")
add_bullet(doc, "Ghost Glass: Backdrop blur (24px/20px), rgba backgrounds, 1px borders, border-radius (24px/16px/12px)")
add_bullet(doc, "Font: Lexend (Google Fonts, weight 100–900)")
add_bullet(doc, "Icons: lucide-react exclusively — no external icon fonts")
add_bullet(doc, "API client: fetchCached() with in-memory TTL cache + request deduplication")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. ZONING ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading("9. Zoning Multi-Tenancy Architecture", level=1)

doc.add_paragraph("The platform is architected for Cell-Based Zoning — enabling deployment across multiple isolated database servers, cloud regions, or on-premise instances. Currently running in standalone mode (Zone 01 baseline).")

doc.add_heading("9.1 Architecture Principles", level=2)
add_bullet(doc, "Cell-Based Isolation: Each Zone is autonomous (Core API + Console UI + PostgreSQL Database).")
add_bullet(doc, "Zone 01 Baseline: Default out-of-the-box deployment runs as Zone 01 (PLATFORM_MODE='standalone').")
add_bullet(doc, "Global Registry: sails_global_master database maps tenant_id → zone_id → zone_api_url.")
add_bullet(doc, "Super Admin War Room: Single-pane dashboard collecting telemetry from all Zones via GET /api/zone/health.")

doc.add_heading("9.2 Global Control Plane Schema", level=2)
doc.add_paragraph("Defined in packages/core/prisma/global.prisma (75 lines, 3 models in global schema):")
add_bullet(doc, "GlobalZone (zones): zone registry — id, name, apiUrl, region, maxTenants, currentTenants, status")
add_bullet(doc, "GlobalTenant (tenants): tenant routing — id, name, slug, domain, zoneId, status (ACTIVE/MIGRATING/SUSPENDED)")
add_bullet(doc, "ZoneHealthMetric (zone_health_metrics): War Room telemetry")
add_bullet(doc, "Status: Schema defined but NOT deployed — reserved for future Phase 6A implementation.")

doc.add_heading("9.3 Zone Telemetry API", level=2)
doc.add_paragraph("GET /api/zone/health (secured via ZONE_SECRET_KEY) returns:")
add_code_block(doc,
    '{\n'
    '  "zoneId": "zone-us-01",\n'
    '  "status": "healthy",\n'
    '  "memoryUsageMB": 184,\n'
    '  "activeDbConnections": 12,\n'
    '  "tenantCount": 42,\n'
    '  "errorCount15m": 0,\n'
    '  "uptimeSeconds": 86400\n'
    '}'
)

doc.add_heading("9.4 TenantConnectionManager", level=2)
doc.add_paragraph("Located at packages/core/src/lib/TenantConnectionManager.ts. When PLATFORM_MODE='zoned':")
add_bullet(doc, "Maintains a Map<string, Pool> cache for per-tenant database connections")
add_bullet(doc, "Resolves DSN dynamically per request based on tenant's target zone")
add_bullet(doc, "5-minute idle pool cleanup via setInterval")
add_bullet(doc, "Current strategy: SCHEMA_PER_TENANT (single DB). Future: DATABASE_PER_TENANT")

doc.add_heading("9.5 Zoning Roadmap (Phase 6, Planned)", level=2)
add_styled_table(doc,
    ["Sub-Phase", "Description"],
    [
        ["6A — Global Control Plane", "Deploy sails_global_master registry mapping tenantId → zoneId → zone_api_url"],
        ["6B — Dynamic DSN Pool Router", "TenantConnectionManager dynamically routes queries across isolated Zone databases"],
        ["6C — Zone Telemetry API", "Expose GET /api/zone/health on every Core API container for monitoring"],
        ["6D — Super Admin War Room", "Console plugin for real-time monitoring of all Zones (AWS, Azure, Local)"],
        ["6E — Zero-Downtime Relocation", "bun run cli tenant:relocate to move tenant schemas across Zones with autonumber continuation"],
    ],
    [4, 12]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. DEPLOYMENT & DEVOPS
# ══════════════════════════════════════════════════════════════
doc.add_heading("10. Deployment & DevOps", level=1)

doc.add_heading("10.1 Current State", level=2)
add_bullet(doc, "NO CI/CD pipeline configured (no .github/ directory, no GitHub Actions, no Jenkins).")
add_bullet(doc, "NO production deployment infrastructure (no Terraform, no Kubernetes, no cloud configs).")
add_bullet(doc, "Development only: Docker Compose with 3 containers (db + core + console).")
add_bullet(doc, "Setting up CI/CD and production deployment is the FIRST priority for the Solution Architect.")

doc.add_heading("10.2 Docker Compose Startup Sequence", level=2)
add_code_block(doc,
    "1. docker compose up -d\n"
    "2. sails-db starts → healthcheck pg_isready → healthy\n"
    "3. sails-core starts (depends_on: db healthy)\n"
    "   → bun prisma generate    (regenerates Prisma client types)\n"
    "   → bun run dev            (Next.js dev server on port 3000)\n"
    "4. sails-console starts (no dependency)\n"
    "   → bun run dev -- --host  (Vite dev server on port 5173)"
)
doc.add_paragraph("Note: prisma generate runs on every container start. No migrations run automatically — they must be applied manually or via init scripts.")

doc.add_heading("10.3 Container Management Rules", level=2)
add_bold_para(doc, "CRITICAL: Never use docker restart on the bun dev containers — they can crash-loop.")
add_code_block(doc, "# Correct way to recreate a container:\n"
    "docker rm -f sails-core && docker compose up -d core\n\n"
    "# Verify compose project (prevents volume mismatch):\n"
    "docker inspect sails-db --format '{{.Config.Labels}}' | grep compose.project"
)
doc.add_paragraph("The Docker volume (klaoplatform_pgdata) is EXTERNAL and persists across container rebuilds. Data is safe during container recreation — only the containers themselves are replaced.")

doc.add_heading("10.4 Build Process", level=2)
doc.add_paragraph("Core Dockerfile (packages/core/Dockerfile):")
add_code_block(doc,
    "FROM oven/bun:latest\n"
    "WORKDIR /app\n"
    "# Layer 1: Copy package manifests only (for caching)\n"
    "COPY package.json bun.lock* ./\n"
    "COPY packages/*/package.json ./packages/*/\n"
    "RUN bun install\n"
    "# Layer 2: Copy source + generate Prisma client\n"
    "COPY . .\n"
    "RUN cd packages/core && bun x prisma generate\n"
    "EXPOSE 3000\n"
    "CMD [\"sh\", \"-c\", \"cd packages/core && bun run dev\"]"
)

doc.add_heading("10.5 Production Considerations (NOT IMPLEMENTED)", level=2)
add_bullet(doc, "PgBouncer support: Both db.ts and knex.ts auto-add ?pgbouncer=true&connection_limit=20 when NODE_ENV=production.")
add_bullet(doc, "NEXTAUTH_SECRET MUST be changed to a strong random value in production.")
add_bullet(doc, "GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET must be configured with real Google OAuth credentials.")
add_bullet(doc, "The console Vite dev server is NOT suitable for production. Use vite build + static hosting (Nginx, S3+CloudFront) for production deployment.")
add_bullet(doc, "Next.js production build: bun run build → bun run start (not bun run dev).")
add_bullet(doc, "Database credentials must be rotated — current mysecretpassword is for development only.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. BACKUP & RECOVERY
# ══════════════════════════════════════════════════════════════
doc.add_heading("11. Backup & Recovery Procedures", level=1)

doc.add_heading("11.1 Standard Backup", level=2)
doc.add_paragraph("Always use the provided script — it produces two timestamped files with guaranteed replay safety:")
add_code_block(doc,
    "./scripts/backup-db.sh\n\n"
    "# Produces:\n"
    "#  backups/sails_schema_YYYYMMDD_HHMMSS.sql  (structure: schemas, tables, indexes, FKs, RLS policies)\n"
    "#  backups/sails_data_YYYYMMDD_HHMMSS.sql    (data: COPY statements, trigger guards for self-referencing FKs)"
)
doc.add_paragraph("Key details:")
add_bullet(doc, "Uses pg_dump -U postgres inside the sails-db container")
add_bullet(doc, "Strips \\restrict/\\unrestrict markers that psql rejects")
add_bullet(doc, "No --create/--clean flags (avoids DROP DATABASE in output)")
add_bullet(doc, "Data dump uses --disable-triggers for self-referencing FK safety")

doc.add_heading("11.2 Restore Procedure", level=2)
add_code_block(doc,
    "docker stop sails-core sails-console\n\n"
    "# 1. Drop existing schemas\n"
    "docker exec sails-db psql -U postgres -c \"DROP SCHEMA IF EXISTS core CASCADE;\"\n"
    "docker exec sails-db psql -U postgres -c \"DROP SCHEMA IF EXISTS tenant_sails_default CASCADE;\"\n\n"
    "# 2. Apply schema, then data\n"
    "cat backups/sails_schema_<TIMESTAMP>.sql | docker exec -i sails-db psql -U postgres -v ON_ERROR_STOP=1\n"
    "cat backups/sails_data_<TIMESTAMP>.sql   | docker exec -i sails-db psql -U postgres -v ON_ERROR_STOP=1\n\n"
    "# 3. Sync any drift\n"
    "docker exec sails-core sh -c \"cd packages/core && bun x prisma migrate diff \\\n"
    "  --from-schema-datasource prisma/schema.prisma \\\n"
    "  --to-schema-datamodel prisma/schema.prisma --script\"\n"
    "# (apply emitted ALTER/CREATE statements manually if any)\n\n"
    "# 4. Restart services\n"
    "docker start sails-core sails-console\n\n"
    "# 5. Sign out and sign in (stale JWT is the #1 cause of 'no data' after restore)"
)

doc.add_heading("11.3 Drift Detection (CRITICAL)", level=2)
add_bold_para(doc, "prisma migrate status DOES NOT detect drift after a DB restore — it only checks _prisma_migrations rows, which come with the dump.")
doc.add_paragraph("The ONLY reliable check is:")
add_code_block(doc,
    "docker exec sails-core sh -c \"cd packages/core && bun x prisma migrate diff \\\n"
    "  --from-schema-datasource prisma/schema.prisma \\\n"
    "  --to-schema-datamodel prisma/schema.prisma --script\"\n\n"
    "# Expected output: -- This is an empty migration.\n"
    "# ANY other output = schema drift that must be resolved."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. KNOWN ISSUES & OPERATIONAL PLAYBOOK
# ══════════════════════════════════════════════════════════════
doc.add_heading("12. Known Issues, Golden Rules & Operational Playbook", level=1)

doc.add_heading("12.1 Six Root Causes of 'No Data in UI' (in order of frequency)", level=2)

add_styled_table(doc,
    ["#", "Root Cause", "Symptoms", "Fix"],
    [
        ["1", "Stale JWT session after DB restore", "Sidebar shows mock fallback (Dashboard/CRM). JWT references deleted user/tenant IDs.", "Sign out and sign in. Do NOT patch code."],
        ["2", "Stale DEFAULT_TENANT_ID in .env", "Config API returns 0 apps. .env points at a deleted tenant.", "Run: SELECT id FROM core.tenants; Update .env, restart sails-core."],
        ["3", "Schema drift (DB older than schema.prisma)", "Column/table does not exist errors. migrate status reports 'up to date' but it's not.", "Run: prisma migrate diff. Apply missing DDL manually. Drop conflicting RLS policies first."],
        ["4", "Auto-mutation in runtime GET handlers", "Navigation suddenly blanks after code change to route.ts.", "Revert the change. Never add writes/seeding to GET handlers. Use migrations or one-off scripts."],
        ["5", "Docker compose project mismatch", "Code edits have no effect. Container name conflicts on docker compose up.", "Check: docker inspect | grep compose.project. Remove stale containers: docker rm -f, then recreate."],
        ["6", "pg_dump full backup cannot run as init script", "DROP DATABASE fails because connected to that database. \\restrict markers cause psql errors.", "Init-restore mount removed from compose. Restore manually via the split backup procedure."],
    ],
    [0.5, 3.5, 5, 7]
)

doc.add_heading("12.2 Golden Rules (Must Not Violate)", level=2)
add_bullet(doc, "1. NEVER put writes, seeding, 'auto-repair', or 'auto-migration' code inside runtime API GET handlers. Use migrations or one-off scripts. (This once blanked the entire navigation.)")
add_bullet(doc, "2. NEVER docker restart the bun dev containers to 'apply changes' — they can crash-loop. Use docker rm -f <name> && docker compose up -d <service>.")
add_bullet(doc, "3. 'No data in the UI' is almost never a code bug. Check order: API response → core logs → browser session (stale JWT) → DEFAULT_TENANT_ID → schema drift.")
add_bullet(doc, "4. ALWAYS use prisma migrate diff (NOT migrate status) to verify schema sync after any restore.")
add_bullet(doc, "5. ALWAYS check compose project name and volume mounts before concluding code isn't reloading.")
add_bullet(doc, "6. NEVER modify RLS policies without checking every column referenced in joined tables — new columns can make unqualified refs ambiguous.")

doc.add_heading("12.3 Diagnosis Checklist", level=2)
add_code_block(doc,
    "# 1. Containers healthy and from the right project?\n"
    "docker ps --format \"table {{.Names}}\\t{{.Status}}\"\n"
    "docker inspect sails-db --format '{{.Config.Labels}}' | grep compose.project\n\n"
    "# 2. What does the API actually return?\n"
    "curl -s http://localhost:3000/api/console/config | python3 -m json.tool | head -20\n\n"
    "# 3. What do the core logs say?\n"
    "docker logs sails-core --since 5m 2>&1 | grep -E \"CONFIG|error|Error\" | tail -10\n\n"
    "# 4. Is there schema drift?\n"
    "docker exec sails-core sh -c \"cd packages/core && bun x prisma migrate diff \\\n"
    "  --from-schema-datasource prisma/schema.prisma --to-schema-datamodel prisma/schema.prisma --script\"\n\n"
    "# 5. Does .env point at a live tenant?\n"
    "docker exec sails-db psql -U postgres -c \"SELECT id, name, schema_name FROM core.tenants;\"\n"
    "grep DEFAULT_TENANT_ID packages/core/.env"
)

doc.add_heading("12.4 Common Gotchas", level=2)
add_bullet(doc, "RLS Policy Gotcha: object_permissions has read_scope/modify_scope enums now. Older dumps use view_all_data/modify_all_data booleans. Adding tenant_id to object_permissions makes unqualified tenant_id references ambiguous in RLS policies — always qualify as 't.tenant_id'.")
add_bullet(doc, "Navigation Sync: Creating a dynamic table is NOT enough to make it appear in the UI. You MUST create a corresponding ConsoleMenu entry linked to a ConsoleApp. See docs/CREATE_APP_NAV.md.")
add_bullet(doc, "System Fields: Metadata tables (core.tables, core.fields) MUST have an is_system boolean column to prevent runtime 500 errors.")
add_bullet(doc, "Admin Password Reset:")
add_code_block(doc,
    "docker exec sails-core sh -c \"cd /app/packages/core && bun -e \\\"\n"
    "import bcrypt from 'bcryptjs'; console.log(await bcrypt.hash('Welcome2Ignite', 12));\\\"\"\n"
    "docker exec sails-db psql -U postgres -c \\\n"
    "  \"UPDATE core.users SET password = '<hash>' WHERE email = 'admin@klao.app';\""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. ROADMAP & TECHNICAL DEBT
# ══════════════════════════════════════════════════════════════
doc.add_heading("13. Roadmap & Technical Debt", level=1)

doc.add_heading("13.1 Completed Phases", level=2)
add_styled_table(doc,
    ["Phase", "Scope", "Status"],
    [
        ["Phase 1", "Foundation & MVP: Core engine, metadata CRUD, field registry", "100%"],
        ["Phase 2", "Dynamic Validation: CHECK constraints, Zod generator, audit trails", "70%"],
        ["Phase 3", "Dynamic UI: Console SPA, DB-driven nav, app switcher, lazy loading", "100%"],
        ["Phase 4", "Multi-Tenancy: Schema-per-tenant, RLS, provisioning, auth bridge", "100%"],
        ["Phase 5", "Verification & Polish: Test suites, Bun migration, Docker dev env", "90%"],
        ["Phase 6", "Advanced Metadata: Capabilities, role-filtered nav, permission UI", "100%"],
        ["Phase 7", "Internal Ops: N:M teams, team ownership, hierarchical RBAC", "100%"],
        ["IAM Phase A", "JWT session, RLS, RBAC, audit logs, Core-Console auth bridge", "Complete"],
        ["IAM Phase B", "Google Workspace OAuth, SSO discovery, domain mapping", "In Progress"],
    ],
    [3, 10, 3]
)

doc.add_heading("13.2 Pending Phases", level=2)
add_styled_table(doc,
    ["Phase", "Scope", "Priority"],
    [
        ["Phase 8", "Internal Modules: UserManager, ObjectManager, Staff Profile", "Current (In Progress)"],
        ["Phase 9", "Staff Profile Refinement: Extended user fields, activation toggle", "Next"],
        ["Phase 10", "Google Workspace Integration: Domain mapping, JIT provisioning, Workspace Settings UI", "Next"],
        ["Phase 2 (remaining)", "Complex validation rules (Regex, Range, Enum plugins)", "Medium"],
        ["Phase 5 (remaining)", "Performance benchmarking for DDL speed", "Low"],
        ["IAM Phase C", "Field-Level Security (FLS) for sensitive data", "Pending"],
        ["IAM Phase D", "Dedicated modules: Timesheets, Case Management", "Pending"],
        ["Phase 2 (Internal Ops)", "Field-Level Security, Manual Sharing, Mass Export guards", "Pending"],
        ["Phase 3 (Projects)", "Criteria-Based Sharing Rules, Auto Timesheet generation", "Pending"],
        ["Phase 4 (PWA)", "Offline-first: Service Workers, Dexie.js, SyncQueue, LWW conflict resolution", "Pending"],
        ["Phase 5 (Google)", "Workspace domain routing, JIT provisioning (already partially implemented)", "Pending"],
        ["Phase 6 (Zoning)", "Cell-Based Zoning (6A–6E): Global Control Plane, Dynamic DSN, War Room", "Planned"],
    ],
    [3, 10, 3]
)

doc.add_heading("13.3 Technical Debt & Missing Infrastructure", level=2)
add_bullet(doc, "NO CI/CD pipeline: No GitHub Actions, no automated testing, no deployment automation.")
add_bullet(doc, "NO production deployment configuration: No Terraform, no Kubernetes manifests, no cloud provider setup.")
add_bullet(doc, "NO monitoring/alerting: No APM, no log aggregation, no uptime monitoring configured.")
add_bullet(doc, "NO staging environment: Only development (Docker Compose) and concept of production.")
add_bullet(doc, "Google OAuth credentials are placeholders (your-google-client-id).")
add_bullet(doc, "NEXTAUTH_SECRET is hardcoded for development (development-secret-only).")
add_bullet(doc, "Database credentials are hardcoded (mysecretpassword).")
add_bullet(doc, "TypeScript strict mode is OFF in core package (strict: false in tsconfig.json).")
add_bullet(doc, "Console production build not validated — only dev server tested.")
add_bullet(doc, "No automated database migration pipeline for production deployments.")
add_bullet(doc, "Schema drift handling is manual (prisma migrate diff + apply by hand).")
add_bullet(doc, "Seed data is minimal — new tenants get standard models but may need domain-specific seeding.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX A: KEY FILE REFERENCE
# ══════════════════════════════════════════════════════════════
doc.add_heading("Appendix A: Key File Reference", level=1)

doc.add_heading("A.1 Core Engine Files", level=2)
add_styled_table(doc,
    ["File", "Role"],
    [
        ["packages/core/src/core/engine/AlchemaCore.ts", "DDL engine — CREATE/ALTER TABLE, RLS policies, triggers, indexes"],
        ["packages/core/src/core/engine/QueryLayer.ts", "Secure DML — RBAC + RLS + Audit, all atomic. THE only way to read/write data."],
        ["packages/core/src/core/engine/AccessGuard.ts", "Object-level RBAC + system capabilities. SUPER_ADMIN fast-path bypass."],
        ["packages/core/src/core/engine/TransactionContext.ts", "Injects SET LOCAL RLS context per transaction (app.current_user_id/tenant_id/team_id)"],
        ["packages/core/src/core/engine/ConnectionManager.ts", "DB pool abstraction (schema-per-tenant → future DB-per-tenant)"],
        ["packages/core/src/core/engine/SchemaLogger.ts", "Fire-and-forget audit logging (data, system, DDL)"],
        ["packages/core/src/core/registry/FieldRegistry.ts", "Singleton field type plugin registry"],
        ["packages/core/src/core/registry/FieldTypePlugin.ts", "Plugin interface: PG column def + Zod schema + React UI components"],
    ],
    [8, 8]
)

doc.add_heading("A.2 Service & Auth Files", level=2)
add_styled_table(doc,
    ["File", "Role"],
    [
        ["packages/core/src/services/TenantProvisioner.ts", "Full tenant bootstrap: schema + user + apps + data models"],
        ["packages/core/src/services/TranslatorLayer.ts", "Metadata ↔ DDL translation bridge"],
        ["packages/core/src/lib/auth/authOptions.ts", "NextAuth config: JWT strategy, Credentials + Google providers, JIT provisioning"],
        ["packages/core/src/lib/auth/session.ts", "Session resolution: getAppSession(), requireSession(), requireAdmin()"],
        ["packages/core/src/lib/db.ts", "PrismaClient singleton"],
        ["packages/core/src/lib/knex.ts", "Raw pg Pool for dynamic queries"],
        ["packages/core/src/lib/zodGenerator.ts", "Dynamic Zod schema from field metadata"],
        ["packages/core/src/lib/TenantConnectionManager.ts", "Zone-aware connection pool manager"],
        ["packages/core/src/lib/configCache.ts", "In-memory config cache (30s TTL)"],
    ],
    [8, 8]
)

doc.add_heading("A.3 Console Key Files", level=2)
add_styled_table(doc,
    ["File", "Role"],
    [
        ["packages/console/src/App.tsx", "Root component — providers + routing"],
        ["packages/console/src/contexts/ConsoleContext.tsx", "Apps, menus, widgets, header state (fetched from /api/console/config)"],
        ["packages/console/src/contexts/AuthContext.tsx", "User session + login/logout"],
        ["packages/console/src/contexts/ThemeContext.tsx", "Dynamic theme/branding from company profile"],
        ["packages/console/src/features/admin/registry.tsx", "AdminPluginRegistry — 18 lazy-loaded components"],
        ["packages/console/src/features/controls/FieldControlRegistry.ts", "20 built-in field controls"],
        ["packages/console/src/features/actions/ActionRegistry.ts", "List/detail toolbar actions"],
        ["packages/console/src/features/widgets/registry.tsx", "Widget bar component registry"],
        ["packages/console/src/pages/DynamicTablePage.tsx", "Universal list view renderer"],
        ["packages/console/src/pages/DynamicDetailPage.tsx", "Universal record detail/form renderer"],
        ["packages/console/src/api/client.ts", "HTTP client with cache + deduplication"],
        ["packages/console/vite.config.ts", "Vite config: dev proxy /api/* → core:3000"],
    ],
    [8, 8]
)

doc.add_heading("A.4 Configuration & Infrastructure", level=2)
add_styled_table(doc,
    ["File", "Role"],
    [
        ["docker-compose.yml (root)", "Primary: 3 services (db + core + console), external volume"],
        ["packages/core/.env", "DATABASE_URL + DEFAULT_TENANT_ID"],
        ["packages/core/prisma/schema.prisma", "Core schema — 21 models in core schema"],
        ["packages/core/prisma/global.prisma", "Global Control Plane schema — 3 models (planned)"],
        ["packages/core/prisma/migrations/", "12 applied migration directories"],
        ["scripts/backup-db.sh", "Full DB backup (schema-only + data-only split dumps)"],
        ["scripts/migrate-system-names.sql", "One-off snake_case renaming migration"],
    ],
    [8, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX B: KEY COMMANDS
# ══════════════════════════════════════════════════════════════
doc.add_heading("Appendix B: Key Commands Reference", level=1)

add_styled_table(doc,
    ["Purpose", "Command"],
    [
        ["Start all services", "docker compose up -d"],
        ["Rebuild and start", "docker compose up -d --build"],
        ["View Core API logs", "docker logs sails-core --since 5m"],
        ["View Console logs", "docker logs sails-console --since 5m"],
        ["SQL shell", "docker exec sails-db psql -U postgres"],
        ["Prisma generate", "docker exec sails-core sh -c \"cd packages/core && bun x prisma generate\""],
        ["Run Prisma migration", "docker exec sails-core sh -c \"cd packages/core && bun x prisma migrate dev --name <name>\""],
        ["Check schema drift", "docker exec sails-core sh -c \"cd packages/core && bun x prisma migrate diff --from-schema-datasource prisma/schema.prisma --to-schema-datamodel prisma/schema.prisma --script\""],
        ["List tenants", "docker exec sails-core sh -c \"cd packages/core && bun run cli tenant:list\""],
        ["Provision tenant", "docker exec sails-core sh -c \"cd packages/core && bun run cli tenant:create \\\"Name\\\" admin@email.com\""],
        ["Sync apps to all tenants", "docker exec sails-core sh -c \"cd packages/core && bun run cli sync-all-tenants\""],
        ["Clean orphaned schemas", "docker exec sails-core sh -c \"cd packages/core && bun run cli db:clean\""],
        ["Check metadata consistency", "docker exec sails-core sh -c \"cd packages/core && bun run cli db:check\""],
        ["Run core tests", "docker exec sails-core bun run test"],
        ["Run security tests", "docker exec sails-core bun run packages/core/tests/test-security.ts"],
        ["Full DB backup", "./scripts/backup-db.sh"],
        ["Reset platform", "docker exec sails-core bun run platform:reset"],
        ["Recreate container (safe)", "docker rm -f sails-core && docker compose up -d core"],
        ["Verify compose project", "docker inspect sails-db --format '{{.Config.Labels}}' | grep compose.project"],
        ["Reset admin password", "See KB_UNLOADED_CONFIG.md §4 or Section 12.4 above"],
        ["Console build (prod)", "cd packages/console && bun run build"],
    ],
    [5, 11]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX C: ENV VAR REFERENCE
# ══════════════════════════════════════════════════════════════
doc.add_heading("Appendix C: Environment Variables Reference", level=1)

add_styled_table(doc,
    ["Variable", "Required", "Default", "Description"],
    [
        ["DATABASE_URL", "Yes", "—", "Primary DB connection string. In-container uses 'db' hostname."],
        ["DEFAULT_TENANT_ID", "Yes", "—", "Fallback tenant ID for unauthenticated config requests."],
        ["NEXTAUTH_SECRET", "Yes", "—", "JWT signing secret. MUST be changed for production."],
        ["NEXTAUTH_URL", "Yes", "—", "Base URL for NextAuth callbacks. Points to console in dev (localhost:5173)."],
        ["GOOGLE_CLIENT_ID", "For OAuth", "placeholder", "Google OAuth client ID."],
        ["GOOGLE_CLIENT_SECRET", "For OAuth", "placeholder", "Google OAuth client secret."],
        ["VITE_CORE_URL", "Yes (console)", "http://localhost:3000", "Console → Core API URL."],
        ["PLATFORM_MODE", "No", "standalone", "When 'zoned', enables multi-database routing via TenantConnectionManager."],
        ["ZONE_ID", "No", "zone-01", "Zone identifier in zoned mode."],
        ["ZONE_SECRET_KEY", "For zoning", "—", "Secures GET /api/zone/health endpoint."],
        ["LOG_DATABASE_URL", "No", "—", "Separate database for audit logs. Falls back to DATABASE_URL."],
        ["LOG_SCHEMA", "No", "core", "Schema for audit log tables."],
        ["TEST_SESSION_JSON", "No", "—", "JSON session override for CLI/integration tests (no browser needed)."],
        ["GLOBAL_DATABASE_URL", "For zoning", "—", "Connection string for the sails_global_master database."],
    ],
    [3.5, 2, 4, 6.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX D: PRISMA MODEL INVENTORY
# ══════════════════════════════════════════════════════════════
doc.add_heading("Appendix D: Prisma Model Inventory", level=1)

doc.add_heading("D.1 Core Schema Models (21 models — schema.prisma)", level=2)
add_styled_table(doc,
    ["Model", "Table", "Key Fields"],
    [
        ["Tenant", "tenants", "id, name, schemaName (unique), createdAt, updatedAt"],
        ["TableDefinition", "tables", "id, tenantId, name, tableName, description, isSystem"],
        ["FieldDefinition", "fields", "id, tableId, name, fieldName, physicalType, logicalType, config (JSONB), isRequired, isSystem, defaultValue, description"],
        ["ValidationRule", "validation_rules", "id, tableId, fieldId?, ruleType (min/max/regex/enum), ruleDefinition, errorMessage"],
        ["TableLayout", "table_layouts", "id, tableId?, tenantId, layoutType, viewType, name, systemName (unique), config (JSONB), publishedConfig (JSONB), status (draft/active), isDefault"],
        ["ConsoleApp", "console_apps", "id, tenantId, name, slug?, description?, icon, order, isSystem, requiredCapability, widgetBarEnabled"],
        ["ConsoleMenu", "console_menus", "id, appId?, parentId?, tenantId, label, icon, path, actionType, componentKey?, dataModelId?, listViewId?, order, isSystem, requiredCapability"],
        ["ConsoleWidget", "console_widgets", "id, tenantId?, appId?, label, icon, componentKey?, openIn, config (JSONB), order, enabled, isSystem, requiredCapability"],
        ["Team", "teams", "id, tenantId, name, isSystemAdmin, parentId? (self-ref), createdAt"],
        ["ObjectPermission", "object_permissions", "id, teamId?, positionId?, userId?, objectName, canCreate, canDelete, readScope (AccessScope enum), modifyScope (AccessScope enum)"],
        ["SystemPermission", "system_permissions", "id, teamId, capability, grantedAt"],
        ["User", "users", "id, tenantId, email (unique per tenant), name?, image?, role (SUPER_ADMIN/TENANT_ADMIN/ADMIN/MEMBER), password?, isActive, emailVerified, googleId?, googleDomain?, phone?, title?, metadata (JSONB), lastLoginAt"],
        ["UserTeam", "user_teams", "userId + teamId (composite PK), isLeader, joinedAt"],
        ["Account", "accounts", "id, userId, type, provider, providerAccountId (composite unique)"],
        ["Session", "sessions", "id, userId, sessionToken (unique), expires (NextAuth standard, unused with JWT)"],
        ["VerificationToken", "verification_tokens", "identifier + token (composite PK)"],
        ["Position", "positions", "id, tenantId, name, prefix, description"],
        ["PositionSlot", "position_slots", "id, positionId, userId, slotNumber, isActive, assignedAt"],
        ["TeamPosition", "team_positions", "teamId + positionId (composite PK)"],
        ["CompanyProfile", "company_profiles", "id, tenantId (unique), companyName, logo?, website?, contact (JSONB), address (JSONB), financial (JSONB), security (JSONB), localization (JSONB), theme (JSONB)"],
        ["DataAuditLog", "data_audit_logs", "id, tenantId, userId?, action (CREATE/UPDATE/DELETE), objectName, recordId?, oldValues (JSONB), newValues (JSONB), timestamp"],
        ["SystemEventLog", "system_event_logs", "id, tenantId, userId?, eventType, eventData (JSONB), timestamp"],
        ["DdlLog", "ddl_logs", "id, tenantId, userId?, ddlStatement, tableName?, timestamp"],
    ],
    [3, 3, 10]
)

doc.add_heading("D.2 Global Schema Models (3 models — global.prisma, PLANNED)", level=2)
add_styled_table(doc,
    ["Model", "Table", "Key Fields"],
    [
        ["GlobalZone", "zones", "id, name, apiUrl, region, maxTenants, currentTenants, status (ZoneStatus enum)"],
        ["GlobalTenant", "tenants", "id, name, slug (unique), domain?, zoneId, status (TenantStatus enum), createdAt"],
        ["ZoneHealthMetric", "zone_health_metrics", "id, zoneId, status, memoryUsageMB, activeDbConnections, tenantCount, errorCount15m, uptimeSeconds, timestamp"],
    ],
    [3, 3, 10]
)

doc.add_heading("D.3 Enums", level=2)
add_styled_table(doc,
    ["Enum", "Schema", "Values"],
    [
        ["AccessScope", "core", "NONE, OWNER, TEAM, HIERARCHY, ALL"],
        ["TenantStatus", "global (planned)", "ACTIVE, MIGRATING, SUSPENDED"],
        ["ZoneStatus", "global (planned)", "HEALTHY, DEGRADED, CRITICAL, MAINTENANCE"],
    ],
    [3, 4, 9]
)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
output_path = "/Users/asana/Repo/Sails Platform/SAILS_Platform_Technical_Handover.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
